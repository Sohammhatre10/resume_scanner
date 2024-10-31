import json
import os
import sys
import docx
from groq import Groq

# Initialize Groq client
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

def create_resume():
    """
    Create a resume JSON structure based on user input.

    Returns:
    dict: Structured resume information in JSON format.
    """
    resume_data = {
        "name": input("Enter your name: "),
        "phone_number": input("Enter your phone number: "),
        "email": input("Enter your email: "),
        "education": [],
        "skills": [],
        "experience": [],
        "certifications": []
    }
    
    # Collect education details
    while True:
        edu = input("Enter your education details (or 'done' to finish): ")
        if edu.lower() == 'done':
            break
        resume_data["education"].append(edu)
    
    # Collect skills
    while True:
        skill = input("Enter a skill (or 'done' to finish): ")
        if skill.lower() == 'done':
            break
        resume_data["skills"].append(skill)
    
    # Collect experience
    while True:
        exp = input("Enter an experience detail (or 'done' to finish): ")
        if exp.lower() == 'done':
            break
        resume_data["experience"].append(exp)
    
    # Collect certifications
    while True:
        cert = input("Enter a certification (or 'done' to finish): ")
        if cert.lower() == 'done':
            break
        resume_data["certifications"].append(cert)
    
    # Optionally enhance resume sections using Groq if API key is provided
    if client.api_key:
        print("Enhancing resume sections using Groq API...")
        try:
            enhanced_resume = enhance_resume_with_groq(resume_data)
            if enhanced_resume:
                resume_data = enhanced_resume
        except Exception as e:
            print(f"Error enhancing resume with Groq: {str(e)}")
    
    return resume_data

def enhance_resume_with_groq(data):
    """
    Enhance resume sections using the Groq API.

    Args:
    data (dict): Resume data to enhance.

    Returns:
    dict: Enhanced resume data.
    """
    response = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": f"Enhance the following resume details for clarity and impact:\n\n{json.dumps(data)}"
            }
        ],
        model="llama-3.2-90b-vision-preview"
    )
    enhanced_text = response.choices[0].message.content

    # Parse JSON output from the response
    try:
        enhanced_data = json.loads(enhanced_text)
        return enhanced_data
    except json.JSONDecodeError:
        print("Failed to parse Groq response as JSON.")
        return data

def save_json_to_file(data, output_file):
    """
    Save JSON data to a file.

    Args:
    data (dict): JSON data to save.
    output_file (str): Path to the output file.
    """
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"JSON resume data has been saved to {output_file}")

def save_resume_to_docx(data, output_file):
    """
    Save resume data to a DOCX file.

    Args:
    data (dict): Resume data in dictionary format.
    output_file (str): Path to the output DOCX file.
    """
    doc = docx.Document()
    doc.add_heading(data["name"], 0)
    doc.add_paragraph(data["phone_number"])
    doc.add_paragraph(data["email"])
    
    # Education Section
    doc.add_heading("Education", level=1)
    for edu in data["education"]:
        doc.add_paragraph(edu, style='List Bullet')
    
    # Skills Section
    doc.add_heading("Skills", level=1)
    for skill in data["skills"]:
        doc.add_paragraph(skill, style='List Bullet')
    
    # Experience Section
    doc.add_heading("Experience", level=1)
    for exp in data["experience"]:
        doc.add_paragraph(exp, style='List Bullet')
    
    # Certifications Section
    doc.add_heading("Certifications", level=1)
    for cert in data["certifications"]:
        doc.add_paragraph(cert, style='List Bullet')
    
    doc.save(output_file)
    print(f"DOCX resume has been saved to {output_file}")

if __name__ == "__main__":
    resume_data = create_resume()
    
    # Save as JSON
    json_output_file = "resume.json"
    save_json_to_file(resume_data, json_output_file)
    
    # Save as DOCX
    docx_output_file = "resume.docx"
    save_resume_to_docx(resume_data, docx_output_file)
    
    print("Resume creation complete.")
